import os
os.environ['OPENAI_API_KEY'] = ''


from pathlib import Path
from langchain_core.prompts import ChatPromptTemplate, PromptTemplate
from langchain_openai import ChatOpenAI
import argparse
from docx import Document
from docx.shared import Inches
import requests
from io import BytesIO

# Set your OpenAI API key
class DocumentGenerator:
    def __init__(self, output_file='data/document.docx'):
        self.output_file = Path(output_file)
        self.llm = ChatOpenAI(api_key=os.environ['OPENAI_API_KEY'], model_name="gpt-3.5-turbo")

    def generate_sections(self, transcript_text):
        """Generate sections in French using a single prompt"""

        # Define the prompt
        prompt = ChatPromptTemplate([
            ("human", "Écrivez un résumé concis, une segmentation thématique, et les points clés du texte suivant en français : {context}")
        ])

        # Format the prompt
        formatted_prompt = prompt.format(context=transcript_text)

        # Generate the response
        response = self.llm(formatted_prompt)

        return response.content

    def process_transcript(self, transcript_path):
        """Process a single transcript and generate a Word document"""
        document = Document()

        transcript_file = Path(transcript_path)
        if transcript_file.exists():
            # Add title
            document.add_heading(transcript_file.stem, level=1)

            # Download and add image from URL
            image_url = 'https://coda.newjobs.com/api/imagesproxy/ms/clu/xnov/xnovagenfrx/branding/165159/Novagen-Conseil-logo.png'
            response = requests.get(image_url)
            image_stream = BytesIO(response.content)
            document.add_picture(image_stream, width=Inches(2.0))

            with open(transcript_file, 'r', encoding='utf-8') as file:
                transcript_text = file.read()
                sections = self.generate_sections(transcript_text)
                document.add_paragraph(str(sections))

            # Save the generated document
            document.save(self.output_file)

            print(f"Document généré : {self.output_file}")
        else:
            print(f"Erreur : Le fichier {transcript_path} n'existe pas.")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Generate a document from a transcript file.')
    parser.add_argument('transcript_path', help='Path to the transcript file')
    args = parser.parse_args()

    generator = DocumentGenerator()
    generator.process_transcript(args.transcript_path)
